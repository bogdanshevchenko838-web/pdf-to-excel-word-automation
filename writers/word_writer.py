from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import OUTPUT_DIR


def _set_cell_border(cell):
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")       # толщина линии
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tcPr.append(element)


def _apply_table_grid(table):
    
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell)


def generate_word(data: dict) -> Path:
    
    invoice_id = data.get("invoice_number") or "output"

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -------------------- TITLE --------------------
    title = doc.add_heading(f"Invoice {invoice_id}", level=1)
    for run in title.runs:
        run.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -------------------- TABLE 1: FIELD / VALUE --------------------
    t1 = doc.add_table(rows=1, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = False
    t1.columns[0].width = Inches(2.0)
    t1.columns[1].width = Inches(4.0)

    hdr = t1.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for c in hdr:
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.bold = True

    header_rows = [
        ("Invoice Number", data.get("invoice_number", "")),
        ("Order Number", data.get("order_number", "")),
        ("Invoice Date", data.get("invoice_date", "")),
        ("Due Date", data.get("due_date", "")),
    ]
    for label, val in header_rows:
        cells = t1.add_row().cells
        cells[0].text = label
        cells[1].text = str(val)

    _apply_table_grid(t1)
    doc.add_paragraph()

    # -------------------- TABLE 2: FROM / TO --------------------
    t2 = doc.add_table(rows=2, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False
    t2.columns[0].width = Inches(3.0)
    t2.columns[1].width = Inches(3.0)

    t2.rows[0].cells[0].text = "From"
    t2.rows[0].cells[1].text = "To"
    for c in t2.rows[0].cells:
        p = c.paragraphs[0]
        for r in p.runs:
            r.bold = True

    t2.rows[1].cells[0].text = data.get("from", "")
    t2.rows[1].cells[1].text = data.get("to", "")

    for cell in t2.rows[1].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _apply_table_grid(t2)
    doc.add_paragraph()

    # -------------------- TABLE 3: ITEMS --------------------
    # 7 колонок: Hrs/Qty | Service | Rate/Price | Adjust | Sub Total | Tax | Total
    t3 = doc.add_table(rows=2, cols=7)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3.autofit = False

    widths = [0.8, 2.2, 1.0, 1.0, 1.0, 0.8, 1.0]
    for col, w in zip(t3.columns, widths):
        col.width = Inches(w)

    headers = ["Hrs/Qty", "Service", "Rate/Price",
               "Adjust", "Sub Total", "Tax", "Total"]
    for i, h in enumerate(headers):
        cell = t3.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True

    r = t3.rows[1].cells
    r[0].text = str(data.get("hrs_qty", ""))
    r[1].text = str(data.get("service", ""))
    r[2].text = str(data.get("rate_price", ""))
    r[3].text = str(data.get("adjust", ""))
    r[4].text = str(data.get("sub_total", ""))
    r[5].text = str(data.get("tax", ""))
    r[6].text = str(data.get("total", ""))

    for p in r[1].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for idx in [0, 2, 3, 4, 5, 6]:
        for p in t3.rows[1].cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _apply_table_grid(t3)
    doc.add_paragraph()

    # -------------------- TABLE 4: TOTALS --------------------
    t4 = doc.add_table(rows=3, cols=2)
    t4.alignment = WD_TABLE_ALIGNMENT.RIGHT
    t4.autofit = False
    t4.columns[0].width = Inches(1.3)
    t4.columns[1].width = Inches(1.3)

    labels = ["Subtotal", "Tax", "Total"]
    values = [
        data.get("sub_total", ""),
        data.get("tax", ""),
        data.get("total", ""),
    ]

    for i in range(3):
        c_label = t4.rows[i].cells[0]
        c_value = t4.rows[i].cells[1]
        c_label.text = labels[i]
        c_value.text = str(values[i])

        p_lbl = c_label.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in p_lbl.runs:
            r.bold = True

        p_val = c_value.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _apply_table_grid(t4)

    # -------------------- SAVE --------------------
    output_dir = Path(OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc_path = output_dir / f"invoice_{invoice_id}.docx"
    doc.save(doc_path)
    print(f"[+] Word file created: {doc_path}")

    return doc_path
